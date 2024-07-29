import asyncio
import io
from asyncio import AbstractEventLoop
from concurrent.futures import Future
from os import listdir
from pathlib import Path
from typing import (
    Callable,
    Optional,
)

from docx import Document
from PIL import Image

from src.configs import Configs
from docx.shared import Mm
from docx.enum.table import WD_TABLE_ALIGNMENT


class CreateDocx(Configs):
    """Формирование документа с таблицей фотографий"""

    PHOTO_DIR_DEFAULT: Path = Path(__file__).parents[1] / Path("photo")
    SUPPORTED_FORMATS: list = ["jpg", "png"]
    OUTPUT: Path = Path(__file__).parents[1] / Path("demo.docx")

    def __init__(
        self,
        loop: AbstractEventLoop,
        callback: Callable[[int, int], None],
        path: str = None,
        col: int = 2,
    ):
        self._completed_files: int = 0
        self._load_test_future: Optional[Future] = None
        self._save_test: Optional[Future] = None
        self._photo_dir = Path(path) if path else self.PHOTO_DIR_DEFAULT
        self._all_file_names = [i for i in listdir(self._photo_dir) if i.split(".")[-1] in self.SUPPORTED_FORMATS]
        self._all_files = len(self._all_file_names)
        self.col = col
        self.row: Optional[int] = None
        self.files: Optional[list[list[Image, str]]] = None
        self._loop = loop
        self._callback = callback
        self._callback_save: Optional[Callable] = None
        self._all_images: dict[str, "CustomImg"] = {}

    def start(self):
        self._load_test_future = asyncio.run_coroutine_threadsafe(self._make_files(), self._loop)

    def cancel(self):
        if self._load_test_future:
            self._loop.call_soon_threadsafe(self._load_test_future.cancel)

    def run_save(self, callback: Callable, url, all_images: dict[str, "CustomImg"]):
        self._all_images = all_images
        self._callback_save = callback
        self._save_test = asyncio.run_coroutine_threadsafe(self._save(url), self._loop)

    async def _save(self, output=None):
        """Сохранение данных в файл"""
        output = self.OUTPUT if output is None else output.name
        doc = Document()
        doc.sections[0].top_margin = Mm(20)
        doc.sections[0].bottom_margin = Mm(20)
        doc.sections[0].left_margin = Mm(25)
        table = doc.add_table(self.row, self.col)
        img_index = 0
        for r in table.rows:
            for cell in r.cells:
                if img_index < len(self.files):
                    p = cell.paragraphs[0]
                    p.alignment = WD_TABLE_ALIGNMENT.CENTER
                    r = p.add_run()
                    img_byte_arr = io.BytesIO()
                    img = self._all_images[f"img_{img_index}"]
                    # img.img_list[0].save(img_byte_arr, format="PNG")
                    self.files[img_index][0].save(img_byte_arr, format="PNG")
                    r.add_picture(img_byte_arr, height=Mm(img.img_height))
                    r.add_text(self._all_file_names[img_index])
                    img_index += 1
        doc.save(output)
        self._callback_save()

    async def _make_files(self):
        """Открытие всех файлов"""
        files = [self._get_files(i) for i in self._all_file_names]
        self.row = self._all_files // self.col + bool(self._all_files % self.col)
        self.files = await asyncio.gather(*files)

    async def _get_files(self, file_name: str):
        """Открытие файла"""
        img = Image.open(self._photo_dir / Path(file_name))
        self._completed_files = self._completed_files + 1
        self._callback(self._completed_files, self._all_files)
        return [img, img.filename.split("/")[-1]]

    async def _create_img_to_docx(self, img_index: int):
        """Подготовка изображения для файла docx"""
        img = self.files[img_index][0]
        wight = self.img_width - self.IMG_PAD * 2
        self.img_height = int(img.size[1] / img.size[0] * wight)
        return img.resize((wight, self.img_height))

